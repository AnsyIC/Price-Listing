from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document

SECTION_RE = re.compile(r"^(?:[一二三四五六七八九十]+、|\d+(?:\.\d+)*[\.、])\s*.+$")
CJK_RE = re.compile(r"[\u4e00-\u9fff]")

NEED_CONFIRMATION_TOKENS = (
    "需确定",
    "待确认",
    "以实际为准",
    "需询价",
    "另计",
)


def _iter_docx_lines(doc: Document) -> Iterable[str]:
    # Paragraph text
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            yield t

    # Table text (some reports may store pricing lines in tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").strip()) for cell in row.cells]
            for cell_text in cells:
                if cell_text:
                    yield cell_text


def read_docx_lines(path: Path) -> List[str]:
    doc = Document(str(path))
    return list(_iter_docx_lines(doc))


def extract_section_headers(lines: List[str], max_headers: int = 25) -> List[str]:
    headers: List[str] = []
    for ln in lines:
        if SECTION_RE.match(ln) and CJK_RE.search(ln):
            headers.append(ln)
            if len(headers) >= max_headers:
                break
    return headers


def _split_name_value(line: str) -> Optional[tuple[str, str]]:
    # Most human lines use Chinese colon, but accept ':' too.
    if "：" in line:
        left, right = line.split("：", 1)
        return left.strip(), right.strip()
    if ":" in line:
        left, right = line.split(":", 1)
        return left.strip(), right.strip()
    return None


def extract_service_lines(lines: List[str], max_items: int = 120) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for ln in lines:
        split = _split_name_value(ln)
        if not split:
            continue
        name, rest = split

        looks_like_money = "元" in rest
        looks_like_note = any(tok in rest for tok in NEED_CONFIRMATION_TOKENS)
        looks_like_total = ("合计" in name) or ("总计" in name)

        if looks_like_money or looks_like_note or looks_like_total:
            items.append({"name": name, "text": rest})
            if len(items) >= max_items:
                break
    return items


def build_case_text_preview(lines: List[str], max_lines: int = 30) -> str:
    return "\n".join(lines[:max_lines])


@dataclass
class CaseCard:
    case_id: str
    source_file: str
    section_headers: List[str]
    service_lines: List[Dict[str, Any]]
    raw_preview: str


def build_case_cards(input_dir: Path) -> List[CaseCard]:
    cards: List[CaseCard] = []
    for path in sorted(input_dir.glob("*.docx")):
        lines = read_docx_lines(path)
        case_id = path.stem
        cards.append(
            CaseCard(
                case_id=case_id,
                source_file=path.name,
                section_headers=extract_section_headers(lines),
                service_lines=extract_service_lines(lines),
                raw_preview=build_case_text_preview(lines),
            )
        )
    return cards


def write_casebank_jsonl(cards: List[CaseCard], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        for card in cards:
            f.write(json.dumps(asdict(card), ensure_ascii=False) + "\n")


def build_casebank(
    input_dir: str | Path = "workflow_data/pricing_reports",
    out_path: str | Path = "workflow_data/casebank.jsonl",
) -> Path:
    input_dir = Path(input_dir)
    out_path = Path(out_path)

    cards = build_case_cards(input_dir)
    write_casebank_jsonl(cards, out_path)
    return out_path


if __name__ == "__main__":
    out = build_casebank()
    print(f"Wrote casebank: {out}")
