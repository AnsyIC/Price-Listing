from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import base64
import datetime

def generate_word(dissected_plan, pricing_report, model_name="gpt-5.2"):
    # Create a new Word document
    doc = Document()

    # Configure fonts for Chinese support
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Also configure Heading styles
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Times New Roman'
            h_style.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Title
    title = doc.add_heading('实验方案报价单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set font for Title manually as it might not inherit from Normal/Heading 1 correctly for eastAsia
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d')}, 生成模型: {model_name}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    CHINESE_NUMBERS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", 
                       "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]

    # Sections
    sections = pricing_report.get('sections', [])
    for i in range(len(sections)):
        section = sections[i]
        header_text = section.get('sectionHeader', '')
        
        # Add Chinese numbering if not present
        prefix = ""
        if i < len(CHINESE_NUMBERS):
            prefix = CHINESE_NUMBERS[i] + "、"
        
        # Check if header already starts with Chinese numbering
        # Simple check: is the first char in CHINESE_NUMBERS?
        # Or just prepend if it doesn't look like it has one.
        # Assuming headers from Agent 1 are just titles like "实验动物"
        final_header = prefix + header_text
        
        doc.add_heading(final_header, level=1)
        
        # Add dissected plan content
        if i < len(dissected_plan):
            content = dissected_plan[i].get('content', '')
            doc.add_paragraph(content)
        
        doc.add_paragraph()  # spacing

        items = section.get('items', [])
        if items:
            # Natural language pricing block
            # Header for pricing
            p_header = doc.add_paragraph()
            run_header = p_header.add_run("报价")
            run_header.bold = True
            
            for item in items:
                name = item.get('name', '')
                unit_price = float(item.get('unitPrice', 0) or 0)
                subtotal = float(item.get('subtotal', 0) or 0)
                qf = item.get('quantityFactors', {})
                is_outsourced = item.get('isOutsourced', False)

                line_text = ""
                
                if unit_price == 0 and subtotal == 0:
                     # Free or Outsourced/Unknown
                     line_text = f"{name}"
                     if not is_outsourced and "免费" not in name:
                         line_text += "（免费）"
                     if is_outsourced:
                         line_text += "（需询价/待确认）"
                else:
                    # Construct unit string: /笼/天
                    key_order = ["只", "天", "次", "笼", "板", "张", "部位", "样本", "抗体", "指标", "个"]
                    unit_suffix = ""
                    sorted_keys = []
                    
                    # Identify keys present in qf
                    present_keys = list(qf.keys())
                    
                    # Sort keys based on key_order
                    for k in key_order:
                        if k in present_keys:
                            unit_suffix += f"/{k}"
                            sorted_keys.append(k)
                    for k in present_keys:
                        if k not in sorted_keys:
                            unit_suffix += f"/{k}"
                            sorted_keys.append(k)
                            
                    # Construct quantity string: 6笼*3天
                    qf_parts = []
                    for k in sorted_keys:
                        qf_parts.append(f"{qf[k]}{k}")
                    qf_str = "*".join(qf_parts)
                    
                    # Format: Name：UnitPrice元/Unit*Quantity=Subtotal元
                    line_text = f"{name}：{unit_price}元{unit_suffix}*{qf_str}={subtotal}元"
                    
                    if is_outsourced:
                         line_text += "（需询价/待确认）"

                p = doc.add_paragraph()
                run = p.add_run(line_text)
                if is_outsourced:
                    run.font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph()  # spacing

    # Final total
    doc.add_heading('总计', level=1)
    total_para = doc.add_paragraph()
    total_run = total_para.add_run(f"总费用: ¥{float(pricing_report.get('totalCost', 0) or 0):.2f}")
    total_run.font.size = Pt(16)
    total_run.font.bold = True

    # Notes
    notes = pricing_report.get('notes', []) or []
    if notes:
        doc.add_heading('备注', level=2)
        for note in notes:
            doc.add_paragraph(str(note), style='List Bullet')

    # Legend
    doc.add_paragraph()
    legend = doc.add_paragraph('标注: ')
    red_run = legend.add_run('红色文字')
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    legend.add_run(' 表示外包/需人工确认项目')

    # Save to binary
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return {
        'filename': 'quotation.docx',
        'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'data': buffer.getvalue()
    }
